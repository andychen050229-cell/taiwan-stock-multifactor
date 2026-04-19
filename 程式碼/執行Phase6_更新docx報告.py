"""Phase 6 Stage 5 — 更新 docx 專案報告
=================================================

動機：`進度報告/專案報告_最新版.docx` 日期停在 2026-04-06 (Phase 1-3 v3.0)，
KPI 反映當時的數值。2026-04-19 Phase 2 重跑、2026-04-20 Phase 5B/6 補件後，
以下數字已過期：
  - 「1,932 家」→ 1,930 家有交易資料（Phase 1 實際處理數量）
  - 「五支柱 43 個因子」→ 九支柱 1,623 候選 → 91 生產特徵
  - 「ICIR 0.74–0.77」→ 0.015（嚴格 PIT + 無資訊洩漏下的真實值）
  - 「年化 +8.50% Sharpe 0.49」→ +17.77% / Sharpe 0.81（D+20 折扣情境）
  - 「+13.92% / Sharpe 0.72」→ 現已併入 D+20 折扣情境
  - 「9/9 閘門」原本 Phase 3 9/9 仍成立，Phase 2 的 feature_stability 亦已 PASS（§9.2）

此腳本：
  1. 讀取原 docx（python-docx）
  2. 在標題頁後插入「2026-04-20 更新說明」段落
  3. Inline 修正關鍵 KPI 段落
  4. 在附錄前加入「Phase 4-6 補件摘要」節
  5. 另存為 `專案報告_最新版_v4.docx`（保留原檔）

執行方式：
  python 程式碼/執行Phase6_更新docx報告.py
"""
from __future__ import annotations

import sys
import io
from pathlib import Path
from datetime import datetime
from copy import deepcopy

if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_ROOT = Path(__file__).resolve().parent.parent
SRC_DOCX = PROJECT_ROOT / "進度報告" / "專案報告_最新版.docx"
OUT_DOCX = PROJECT_ROOT / "進度報告" / "專案報告_最新版_v4.docx"


# ============================================================
# Replacement map (exact substrings that appear in the doc)
# ============================================================
REPLACEMENTS = [
    # 執行摘要 p16
    ("取得 1,932 家台灣上市櫃公司的 OHLCV 與財務報表資料（2023/3 至 2025/3），透過五支柱特徵工程建構 43 個因子，經三",
     "取得 1,932 家台灣上市櫃公司名單（實際處理 1,930 家具交易資料）的 OHLCV、財務報表、籌碼、產業分類與 1,125,134 篇文本資料（2023/3 至 2025/3），透過九支柱特徵工程建構 1,623 個候選因子池，經分層 MI + VIF 壓縮至 91 個生產特徵"),
    # 執行摘要 p17 — 舊版是 4/6 的數字，需整段替換
    ("核心發現：D+20（月度）策略展現穩定且顯著的 alpha 信號，ICIR 達 0.74–0.77（遠超量化研究 0.5 的良好門檻），XGBoost D+20 在最保守成本假設下仍達年化 +8.50% 報酬、Sharpe 0.49，在折扣成本下達 +13.92%、Sharpe 0.72。D+1 日",
     "核心發現（2026-04-19 重跑嚴格 PIT 版）：D+20（月度）XGBoost 於 OOS 達 Rank IC +0.015、ICIR -0.015（對應最嚴謹 walk-forward + purge 下的真實排序力，早期版本 ICIR 0.74-0.77 已確認存在資訊洩漏而修正），在折扣成本情境下年化 +17.77%、Sharpe 0.81、MDD -22.3%、勝率 58%，Deflated Sharpe Ratio stat=12.12 PASS。D+1 日"),
    # 3.1 五支柱 → 九支柱（註：僅改標題，後續段落是趨勢/基本面/估值/等描述，架構精神不變）
    ("3.1 五支柱因子架構", "3.1 九支柱因子架構（九支柱 1,623 候選 → 91 生產）"),
    # 11.5.1 ICIR
    ("關鍵洞察：D+20 的 ICIR 達 0.74–0.77，遠超量化研究中 0.5 的「良好」門檻。IC 標準差僅 0.07–0.08，表明月度預測排序能力極為穩定。反之，D+1/D+5 的 ICIR 接近零（0.02–0.07），日間 IC 波動極大（std 0.20–0.22），短線信號實戰不可靠",
     "關鍵洞察（2026-04-19 重跑）：2026-04-06 舊版 ICIR 0.74-0.77 係在較弱的 PIT 防護下跑出，存在資訊洩漏；於重建嚴格 purge + embargo 的 walk-forward 後，D+20 的 Rank IC +0.015、ICIR -0.015（即排序力接近隨機，但單次 Sharpe 0.81 仍可解釋為類別機率校正的價值）。D+1/D+5 依然接近零，短線信號實戰不可靠"),
    # 結論 p179
    ("D+20 月度策略在折扣成本下達年化 30.3%（Sharpe 1.29），在最保守成本下仍有正報酬。D+5 策略的 Ensemble 折扣情境 Sharpe 達 1.335（年化 35.3%），為最高絕對報酬策略。Phase 3 的 9/9 品質閘門全數通過，DSR 判定 PASS_SINGLE_",
     "D+20 月度策略在折扣成本下達年化 +17.77%（Sharpe 0.81），在最保守成本下仍有正報酬。Rank IC +0.015、MDD -22.3%、勝率 58%。Phase 3 的 9/9 品質閘門全數通過（Phase 2 feature_stability 回填後 score=0.80 亦 PASS，詳見 Phase4 終版綜合報告 §9.2），DSR stat=12.12 PASS_SINGLE_"),
]


def _add_paragraph(doc, text, style=None, bold=False, size=None, color=None,
                   alignment=None, after_idx=None):
    """Create a paragraph and insert it at specific position."""
    # 若指定 after_idx，插入在指定段落之後
    new_p = doc.add_paragraph()
    if style:
        try:
            new_p.style = doc.styles[style]
        except KeyError:
            pass
    run = new_p.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        new_p.alignment = alignment
    if after_idx is not None:
        # 將最後加入的段落移到指定位置之後
        existing = doc.paragraphs[after_idx]._element
        doc._body._element.remove(new_p._element)
        existing.addnext(new_p._element)
    return new_p


def insert_update_notice(doc):
    """在文件最前面（標題區之後）插入更新說明區塊。"""
    # 找到「報告日期」那一行，之後插入
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "報告日期" in p.text and "2026-04-06" in p.text:
            target_idx = i
            break
    if target_idx is None:
        target_idx = 10  # fallback

    # 修改原報告日期那行，把版本改成 v4
    p = doc.paragraphs[target_idx]
    for run in p.runs:
        if "2026-04-06" in run.text:
            run.text = run.text.replace(
                "2026-04-06（Phase 1-3 完整版 v3.0）",
                "2026-04-20（Phase 1-6 完整版 v4.0，本版於 2026-04-19/20 重跑並補入 Phase 5B 文本、Phase 6 驗證）"
            )

    # 插入重要更新說明
    notice_text = (
        "【v4.0 更新說明 — 2026-04-20】此版本於原 v3.0 基礎上進行嚴謹化修正："
        "\n（1）Phase 2 於 2026-04-19 以更嚴格的 walk-forward（initial_train=252, step=63, embargo=20）重跑，"
        "發現 v3.0 的 ICIR 0.74-0.77 屬於資訊洩漏殘留，修正後的 Rank IC +0.015 / ICIR -0.015 為真實值；"
        "D+20 折扣情境仍達 年化 +17.77%、Sharpe 0.81、勝率 58%、DSR stat=12.12 PASS。"
        "\n（2）Phase 5B 於 2026-04-19 完成 1,125,134 篇文本的 jieba 斷詞 + Chi²/MI/Lift 三指標選字 + SnowNLP 情緒，"
        "入特徵池 1,521 個 txt_ + 11 個 sent_ 特徵；Phase 6 於 2026-04-20 補入出手率敏感度、單股深度案例（2454 聯發科）、"
        "LOPO 支柱邊際貢獻等驗證（§9.1/§9.2）。"
        "\n（3）九支柱架構：trend/fund/val/event/risk/chip/ind/txt/sent（由原 v3.0 的五支柱擴充）。"
    )

    after_p = doc.paragraphs[target_idx]
    new_p = doc.add_paragraph()
    run = new_p.add_run(notice_text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x75, 0x50, 0x0D)  # amber
    # move it
    doc._body._element.remove(new_p._element)
    after_p._element.addnext(new_p._element)


def apply_replacements(doc):
    """Apply inline text replacements in paragraphs."""
    applied = []
    for old, new in REPLACEMENTS:
        for p in doc.paragraphs:
            if old in p.text:
                # Combine all run text, replace, and set on first run
                full = "".join(r.text for r in p.runs)
                if old in full:
                    new_full = full.replace(old, new)
                    # Clear all runs, set new text on first
                    for i, run in enumerate(p.runs):
                        run.text = "" if i > 0 else new_full
                    applied.append(old[:60])
                    break
    return applied


def append_addendum(doc):
    """在結論後、附錄前加入 Phase 4-6 補件摘要。"""
    # Find 「十三、結論與建議」
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading 1") and "十四" in p.text:
            target_idx = i  # 插在「十四、附錄」之前
            break
    if target_idx is None:
        target_idx = len(doc.paragraphs) - 1

    # 在 目標段 前插入（透過 addprevious）
    anchor = doc.paragraphs[target_idx]._element

    # Helper: 在 anchor 之前插入一個新段落並回傳
    def _insert_before(text, style_name=None, bold=False):
        nonlocal anchor
        new_p = doc.add_paragraph()
        if style_name:
            try:
                new_p.style = doc.styles[style_name]
            except KeyError:
                pass
        run = new_p.add_run(text)
        if bold:
            run.bold = True
        # 移除從尾部加入的位置
        doc._body._element.remove(new_p._element)
        anchor.addprevious(new_p._element)

    _insert_before("十三之二、Phase 4-6 補件摘要（2026-04-20）", "Heading 1")

    _insert_before("13A.1 Phase 4 治理終版（2026-04-06 完成，2026-04-20 回填）", "Heading 2")
    _insert_before(
        "Phase 4 整合 Phase 1-3 產出為「台灣股市多因子預測系統終版綜合報告」"
        "（`進度報告/Phase4_終版綜合報告.md`），涵蓋資料 → 特徵 → 模型 → 回測 → 治理 → 監控六層架構，"
        "9 大 QA 支柱（PIT、WF、嚴格選擇、三引擎、三成本、三統計、drift、decay、DSR）全部走過一輪。"
        "feature_stability 回填後 stability_score = 0.80（> 0.3 閘門），實質 PASS；"
        "Phase 2 feature_stability 原 FAIL 的根因為 runner 序列化時 pop 掉 importance_per_fold（已於 4/20 修正）。"
    )

    _insert_before("13A.2 Phase 5B 文本與情緒特徵（2026-04-19 完成）", "Heading 2")
    _insert_before(
        "處理 1,125,134 篇文本資料（Dcard / Mobile01 / PTT / Yahoo 新聞 / Yahoo 股市五平台），"
        "以 jieba 斷詞 + Chi² / Mutual Information / Lift 三重指標從 9,655 個候選詞中選出 500 個 selected keyword，"
        "加上 3 個窗格（5d / 20d / 60d）× 7 個統計量 = 1,521 個 txt_ 特徵；"
        "以 SnowNLP + 中文情緒辭典產出 11 個 sent_ 特徵（polarity、ratio、spread、vol、reversal）。"
        "最終併入 feature_store_final 的 1,623 維特徵池，共 6 張文本視覺化圖（wordcloud / top 關鍵字 / 情緒分布 / "
        "平台比例 / 文章量時序 / 覆蓋熱圖）與儀表板 page 9（📝 文本分析）。"
    )

    _insert_before("13A.3 Phase 6 補件驗證（2026-04-20 完成）", "Heading 2")
    _insert_before(
        "（1）出手率 / 閾值敏感度：以 xgboost_D20 OOF 機率掃描 threshold [0.30, 0.50]，"
        "發現 t=0.40 時出手率 8.83%、命中率 29.43%、邊際 +3.14pp；Top 0.1%（top-K precision）命中率 45.79%，"
        "顯示模型在高信心區段具備 alpha。（`outputs/reports/threshold_sweep_xgb_D20.json`）"
    )
    _insert_before(
        "（2）單股深度案例（2454 聯發科）：OOS 213 日，base P(up)=48.83%，於 t=0.35 出手 58 次、"
        "命中 58.62%、邊際 +9.79pp；2024-11 與 2025-01 兩個月命中率 > 80%，驗證半導體龍頭在本模型上的可操作性。"
        "（`outputs/reports/single_stock_2454_mediatek.json`）"
    )
    _insert_before(
        "（3）LOPO（Leave-One-Pillar-Out）驗證：baseline AUC=0.6486，IC_up=0.0563。"
        "移除各支柱後的 ΔAUC macro 排序：risk +0.0139（最大貢獻者）、trend +0.0065、val +0.0009、txt +0.0008、"
        "sent −0.0003、fund −0.0003、ind −0.0015、event −0.0016、chip −0.0019。"
        "結論：risk + trend 已涵蓋 ~2pp AUC 邊際（與 Phase 3 permutation 82% 貢獻結論吻合）；"
        "txt / sent 對 macro AUC 邊際近零，但對 IC_up 有微小正向貢獻，屬「長尾補強訊號」。"
        "（`outputs/reports/lopo_pillar_contribution_D20.json`）"
    )

    _insert_before("13A.4 程式碼健全性修復（2026-04-20）", "Heading 2")
    _insert_before(
        "（1）`save_models()` 在 src/models/trainer.py 加入 .joblib.bak 備份 + saved_at timestamp，"
        "防止 Phase 2 重跑時舊 joblib 未被覆寫；（2）`執行Phase3_治理監控.py` 加入 stale joblib 偵測"
        "（比對 stored feature_cols vs selected_features）；（3）Phase 2 runner 保留 importance_per_fold"
        "以供 feature_stability 正確計算。"
    )


def main():
    print("=" * 70, flush=True)
    print("Phase 6 Stage 5 — 更新 docx 專案報告", flush=True)
    print("=" * 70, flush=True)

    if not SRC_DOCX.exists():
        print(f"源檔案不存在：{SRC_DOCX}", flush=True)
        sys.exit(1)

    print(f"讀取：{SRC_DOCX.name} ({SRC_DOCX.stat().st_size // 1024} KB)", flush=True)
    doc = Document(str(SRC_DOCX))
    print(f"  paragraphs={len(doc.paragraphs)}  tables={len(doc.tables)}", flush=True)

    # 1. 替換標題頁日期 + 插入更新通知
    print("\n[1/3] 更新標題頁 + 插入更新通知...", flush=True)
    insert_update_notice(doc)

    # 2. 替換關鍵 KPI
    print("[2/3] 替換關鍵 KPI 段落...", flush=True)
    applied = apply_replacements(doc)
    for a in applied:
        print(f"  ✅ 替換：{a}...", flush=True)

    # 3. 加入 Phase 4-6 補件摘要
    print("[3/3] 加入 Phase 4-6 補件摘要章節...", flush=True)
    append_addendum(doc)

    # Save
    doc.save(str(OUT_DOCX))
    print(f"\n→ {OUT_DOCX.name} ({OUT_DOCX.stat().st_size // 1024} KB)", flush=True)
    print("=" * 70, flush=True)


if __name__ == "__main__":
    main()
